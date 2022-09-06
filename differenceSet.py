# 读取docx中的文本代码示例
import sys

import docx

# 获取文档对象
file = docx.Document(
    r"总字典路径.docx")
file2 = docx.Document(
    r"需要删除的小字典路径.docx")

# 把所有字加到一个列表中，创建一个字典
total_text = []
for para in file.paragraphs:
    for run_1 in para.runs:
        total_text.extend(run_1.text)

# 挨个字对比
num = 0  # 文档汉字个数
for para_2 in file2.paragraphs:
    # if para_2.text == "标题":  # 跳过这个文件的标题
    #     continue
    for run in para_2.runs:
        if '\u4e00' <= run.text <= '\u9fff':  # 判断是否为汉字
            if run.text in total_text:
                index = total_text.index(run.text)
                # total_text.remove(run.text)
                del total_text[index:index + 3]  # 删除汉字，以及汉字后面的两个空格
                if run.text in total_text:
                    print("总库中含有多个该汉字：" + run.text)
                    print("程序中止，检查后重新尝试。")
                    sys.exit(0)
                print(run.text + "：从总库中移除")
                num += 1

# 写入文件
doc = docx.Document()
temp = []
text_i = 0
for _ in range(len(total_text)):
    if text_i >= len(total_text):
        break
    if ('\u0041' <= total_text[text_i] <= '\u005a') or ('\u0061' <= total_text[text_i] <= '\u007a'):  # 字母占一行
        doc.add_paragraph(total_text[text_i:text_i + 3])  # 根据自己情况修改，我的字典里每个字母后面有两个空格，所以都加上
        text_i += 3
        continue
    else:
        temp.extend(total_text[text_i])
        text_i += 1
        if text_i >= len(total_text):
            doc.add_paragraph(temp)
            temp = []
        elif ('\u0041' <= total_text[text_i] <= '\u005a') or ('\u0061' <= total_text[text_i] <= '\u007a'):
            doc.add_paragraph(temp)
            temp = []

print("共移除" + str(num) + "个汉字")
doc.save(r"保存路径.docx")
print("文件已保存")
